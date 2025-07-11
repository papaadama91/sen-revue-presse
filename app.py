import streamlit as st
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from PIL import Image
import numpy as np
import os
import re
import easyocr
from transformers import pipeline
from collections import defaultdict
from deep_translator import GoogleTranslator
from gtts import gTTS
from moviepy.editor import *
from PIL import Image, ImageDraw, ImageFont

reader = easyocr.Reader(['fr'])

@st.cache_resource
def load_summarizer():
    return pipeline("summarization", model="facebook/bart-large-cnn")

summarizer = load_summarizer()

def extract_text(file):
    ext = os.path.splitext(file.name)[-1].lower()
    if ext == ".pdf":
        reader_pdf = PdfReader(file)
        return "\n".join([page.extract_text() for page in reader_pdf.pages if page.extract_text()])
    elif ext == ".docx":
        doc = DocxDocument(file)
        return "\n".join([p.text for p in doc.paragraphs])
    elif ext == ".txt":
        return file.read().decode("utf-8")
    elif ext in [".jpg", ".jpeg", ".png"]:
        image = Image.open(file).convert("RGB")
        results = reader.readtext(np.array(image), detail=0)
        return "\n".join(results)
    else:
        return "[Format non supporté]"

def extract_metadata(text):
    title = text.strip().split("\n")[0][:80]
    date_match = re.search(r'\d{1,2} \w+ \d{4}', text)
    date = date_match.group() if date_match else "Date inconnue"
    return title, date

def classify_article(text):
    keywords = {
        "Politique": ["président", "gouvernement", "élection"],
        "Économie": ["croissance", "pib", "bourse", "marché", "inflation"],
        "Société": ["éducation", "santé", "social", "famille", "logement"],
        "Culture": ["musique", "cinéma", "livre", "art", "festival"],
        "Sport": ["football", "match", "jo", "ligue", "compétition"]
    }
    for theme, mots in keywords.items():
        if any(mot in text.lower() for mot in mots):
            return theme
    return "Autres"

def summarize_text(text):
    try:
        return summarizer(text[:1024], max_length=100, min_length=30, do_sample=False)[0]['summary_text']
    except:
        return "Résumé non disponible."

def translate_to_wolof(text):
    try:
        return GoogleTranslator(source='fr', target='sn').translate(text)
    except:
        return "[Traduction indisponible]"

def generate_voiceover_script(articles_by_theme, translate=False):
    script = ""
    for theme, articles in articles_by_theme.items():
        script += f"\n🗂️ {theme}\n"
        for art in articles:
            phrase = f"Le {art['date']}, un média local a rapporté : {art['resume']}\n"
            if translate:
                phrase = translate_to_wolof(phrase)
            script += phrase
    return script.strip()

def generate_audio_fr(text, filename="voix_off_fr.mp3"):
    tts = gTTS(text=text, lang="fr")
    tts.save(filename)
    return filename

from docx import Document as DocxDocumentExport
def generate_docx_by_theme(articles_by_theme):
    doc = DocxDocumentExport()
    doc.add_heading("Revue de Presse Thématique", 0)
    for theme, articles in articles_by_theme.items():
        doc.add_heading(f"Thème : {theme}", level=1)
        for art in articles:
            doc.add_heading(art['titre'], level=2)
            doc.add_paragraph(f"Date : {art['date']}")
            doc.add_paragraph("Résumé :")
            doc.add_paragraph(art['resume'])
            doc.add_paragraph("---")
    path = "revue_presse.docx"
    doc.save(path)
    return path

# Interface Streamlit
st.title("📰 Générateur de Revue de Presse Thématique (EasyOCR)")
st.markdown("Charge tes fichiers presse ci-dessous (.pdf, .docx, .txt, .png, .jpg)")

uploaded_files = st.file_uploader("Fichiers :", accept_multiple_files=True, type=["pdf", "docx", "txt", "png", "jpg", "jpeg"])

if uploaded_files:
    articles_by_theme = defaultdict(list)
    with st.spinner("Analyse des fichiers en cours..."):
        for file in uploaded_files:
            text = extract_text(file)
            titre, date = extract_metadata(text)
            theme = classify_article(text)
            resume = summarize_text(text)
            article = {"titre": titre, "date": date, "theme": theme, "resume": resume}
            articles_by_theme[theme].append(article)

    if articles_by_theme:
        docx_path = generate_docx_by_theme(articles_by_theme)
        st.success("✅ Revue de presse générée.")
        with open(docx_path, "rb") as f:
            st.download_button("📥 Télécharger Word", f, file_name="revue_presse.docx")

        st.subheader("🎙️ Script Voix-Off (Français)")
        fr_script = generate_voiceover_script(articles_by_theme)
        st.text_area("Script FR", fr_script, height=300)

        st.subheader("🌍 Script Voix-Off (Wolof)")
        wo_script = generate_voiceover_script(articles_by_theme, translate=True)
        st.text_area("Script WO", wo_script, height=300)

        st.download_button("📤 Télécharger Script FR", fr_script, file_name="script_voixoff_fr.txt")
        st.download_button("📤 Télécharger Script WO", wo_script, file_name="script_voixoff_wolof.txt")

        if st.button("🔉 Générer voix-off FR (MP3)"):
            audio_path = generate_audio_fr(fr_script)
            with open(audio_path, "rb") as audio_file:
                st.audio(audio_file.read(), format="audio/mp3")
                st.download_button("📥 Télécharger MP3", audio_file, file_name="voix_off_fr.mp3")
        if st.button("🎥 Générer vidéo YouTube (MP4)"):
            video_path = generate_news_video_enhanced(fr_script, audio_path)
            with open(video_path, "rb") as vid:
                st.video(vid)
                st.download_button("📽 Télécharger la vidéo", vid, file_name="revue_presse_video.mp4")

def generate_news_video_enhanced(
    script_text,
    audio_path,
    logo_path="logo_senegal.png",     # doit être placé à la racine du projet
    music_path=None,                  # facultatif
    output_path="revue_presse_video.mp4"
):
    # Durée audio
    audio_clip = AudioFileClip(audio_path)
    duration_audio = audio_clip.duration
    width, height = 1280, 720
    bg_color = (10, 10, 40)

    # Créer fond d'écran
    background = Image.new("RGB", (width, height), bg_color)
    draw = ImageDraw.Draw(background)

    # Polices
    title_font = ImageFont.truetype("DejaVuSans-Bold.ttf", 60)
    text_font = ImageFont.truetype("DejaVuSans.ttf", 34)

    # Titre
    draw.text((50, 30), "📰 Revue de Presse - Sénégal", font=title_font, fill="white")

    # Script (80 caractères max par ligne)
    wrapped = "\n".join(script_text[i:i+80] for i in range(0, len(script_text), 80))
    draw.text((50, 150), wrapped, font=text_font, fill="white")

    # Logo
    try:
        logo = Image.open(logo_path).resize((100, 100)).convert("RGBA")
        background.paste(logo, (width - 140, 40), logo)
    except Exception as e:
        print(f"Aucun logo trouvé ({e}), on continue sans.")

    # Enregistrer image
    img_path = "fond_temp.jpg"
    background.save(img_path)

    # Créer clip
    clip = ImageClip(img_path).set_duration(duration_audio)
    clip = clip.set_audio(audio_clip)

    # Ajouter musique douce si fournie
    if music_path and os.path.exists(music_path):
        music = AudioFileClip(music_path).volumex(0.1)
        final_audio = CompositeAudioClip([audio_clip.volumex(1.0), music.set_duration(duration_audio)])
        clip = clip.set_audio(final_audio)

    # Export
    clip.write_videofile(output_path, fps=24)
    return output_path
